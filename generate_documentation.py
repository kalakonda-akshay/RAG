"""
Generates comprehensive Word (.docx) and PDF (.pdf) Technical & Executive Documentation
for the Orbit Offline Multimodal RAG Platform.
"""
import os
import sys

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(PROJECT_DIR, "data", "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)

DOCX_PATH = os.path.join(REPORTS_DIR, "Orbit_Offline_RAG_Platform_Documentation.docx")
PDF_PATH = os.path.join(REPORTS_DIR, "Orbit_Offline_RAG_Platform_Documentation.pdf")

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls
except ImportError:
    docx = None

try:
    import pymupdf as fitz
except ImportError:
    fitz = None


def generate_docx_documentation():
    if docx is None:
        print("python-docx not installed, skipping docx generation.")
        return

    doc = docx.Document()

    # Define Color Palette
    PRIMARY_COLOR = RGBColor(99, 102, 241)     # Indigo #6366F1
    SECONDARY_COLOR = RGBColor(168, 85, 247)   # Purple #A855F7
    TEXT_COLOR = RGBColor(31, 41, 55)         # Dark Gray #1F2937
    LIGHT_BG = "F3F4F6"

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("🧠 ORBIT PLATFORM")
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Offline Multimodal Knowledge Engine & AI Assistant — Technical & Executive Documentation")
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_COLOR
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = SECONDARY_COLOR
        return p

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.bold = True
            r_b.font.size = Pt(11)
            r_b.font.color.rgb = TEXT_COLOR
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT_COLOR
        return p

    # 1. Executive Summary
    add_heading_1("1. Executive Summary & Product Vision")
    add_body("Orbit is an enterprise-grade, 100% offline Multimodal Retrieval-Augmented Generation (RAG) platform designed to provide private, air-gapped document intelligence without transmitting sensitive data to external cloud APIs.")
    add_body("Privacy & Zero Cloud Leakage: ", "Key Value Proposition: ")
    add_body("Operates completely air-gapped on local hardware with zero data transmission to external servers, satisfying strict GDPR, HIPAA, and SOC-2 compliance standards.")
    add_body("Multimodal Document Processing: ", "Broad Format Coverage: ")
    add_body("Native support for PDFs, Word (.docx), PowerPoint (.pptx), Excel/CSV spreadsheets, Images (.png, .jpg, .webp), Audio files (.mp3, .wav), Video files (.mp4), Web Archives (.html), and SQLite databases (.db).")

    # 2. System Architecture & How App Works
    add_heading_1("2. System Architecture & How the App Works")
    add_body("The Orbit platform follows an end-to-end local data lifecycle:")
    
    add_heading_2("A. 3-Column UI Architecture (Orbit Dashboard)")
    add_body("Leftmost sidebar containing user profile, primary reset actions, core/management navigation, active AI role switcher, model selector, and file uploader.", "Column 1 (AppSidebar): ")
    add_body("Central inbox pane rendering indexed document cards, search filters, and date-stamped conversation streams.", "Column 2 (Middle List Pane): ")
    add_body("Interactive chat workspace, voice call turn handler, Mermaid mind maps, topic clusters, SQL query execution, and report generators.", "Column 3 (Right Workspace Pane): ")

    add_heading_2("B. Data Ingestion & Indexing Pipeline")
    add_body("Auto-detects file format via extension routing (`ingestion/router.py`) and passes files to dedicated parsers.", "1. Multi-Format Detection: ")
    add_body("Extracted text is split into overlapping chunks (300 words with 50-word overlap) using Recursive Character Chunking (`core/chunker.py`).", "2. Chunking & Overlap: ")
    add_body("Generates 768-dimensional dense vector embeddings using local `nomic-embed-text` via Ollama (`core/embedder.py`).", "3. Vector Embedding: ")
    add_body("Chunks and embeddings are stored in a persistent local ChromaDB instance with HNSW vector indexing (`core/vectorstore.py`).", "4. Persistent Storage: ")

    add_heading_2("C. Retrieval & Prompt Synthesis Pipeline")
    add_body("Decomposes multi-part queries into targeted sub-questions using LLM reasoning (`core/agentic_rag.py`).", "1. Agentic Decomposition: ")
    add_body("Combines Dense Vector Cosine Similarity with Sparse BM25 Keyword Scoring via Reciprocal Rank Fusion (RRF).", "2. Hybrid Search (BM25 + Dense): ")
    add_body("Combines retrieved chunks, conversation history, and persona prompt into a grounded context window (`core/rag_engine.py`).", "3. Persona Prompting & Memory: ")
    add_body("A Self-Reflective Critic Agent verifies factual consistency and eliminates hallucinations before outputting the final response (`core/self_rag.py`).", "4. Critic Verification (Self-RAG): ")

    # 3. Technical Stack & Rationale
    add_heading_1("3. Complete Technical Stack & Engineering Rationale")
    
    # Table of Tech Stack
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component"
    hdr_cells[1].text = "Technology Selected"
    hdr_cells[2].text = "Engineering Rationale"

    tech_data = [
        ("Frontend GUI", "Streamlit 1.39+ & CSS Glassmorphism", "Provides responsive, rapid UI rendering with full custom CSS control for 3-column layouts."),
        ("Local LLM Engine", "Ollama (llama3.2:3b)", "Ultra-fast sub-second latency with 3B parameter reasoning on commodity CPU hardware."),
        ("Vector Embedding", "Ollama (nomic-embed-text)", "High-performance 768-dim embedding model optimized for technical retrieval."),
        ("Vector Database", "ChromaDB Persistent Client", "Zero-server, local SQLite-backed HNSW vector store requiring no cloud infrastructure."),
        ("Keyword Search", "Rank-BM25 (BM25Okapi)", "Provides exact keyword, acronym, and part-number matching complementing dense search."),
        ("PDF Processing", "PyMuPDF (fitz) + Tesseract OCR", "Lightning-fast PDF parsing with automatic OCR fallback for scanned documents."),
        ("Document Parsing", "python-docx & python-pptx", "Native parsing of Microsoft Office documents, slides, speaker notes, and tables."),
        ("Spreadsheet RAG", "pandas & openpyxl", "Converts CSV and Excel sheets into structured Markdown tables for accurate table RAG."),
        ("Audio Transcription", "faster-whisper (CTranslate2)", "Int8-quantized speech-to-text with VAD filtering running 4x faster than real time."),
        ("Video Processing", "imageio-ffmpeg", "Extracts 16kHz mono audio streams from .mp4/.mkv video files for Whisper transcription."),
        ("Text-to-Speech", "pyttsx3 (SAPI5)", "100% offline native speech synthesis with speed & volume control."),
        ("Bundle Packaging", "PyInstaller 6.22", "Bundles Python runtime, C DLLs, and dependencies into a single 1-click executable."),
    ]

    for row in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. Algorithms & Mathematical Formulations
    add_heading_1("4. Algorithms & Mathematical Formulations")

    add_heading_2("A. Reciprocal Rank Fusion (RRF)")
    add_body("To combine sparse keyword ranks (BM25) and dense vector similarity ranks (Cosine) into a single unified score, Orbit implements Reciprocal Rank Fusion:")
    add_body("RRF_Score(d) = Σ [ 1 / (k + r_m(d)) ]   where k = 60", "RRF Formula: ")
    add_body("RRF prevents high scores in one retriever from overriding another, ensuring balanced hybrid retrieval precision.")

    add_heading_2("B. Okapi BM25 Keyword Scoring")
    add_body("Score(D, Q) = Σ IDF(q_i) * [ f(q_i, D) * (k1 + 1) ] / [ f(q_i, D) + k1 * (1 - b + b * (|D| / avgdl)) ]", "BM25 Formula: ")
    add_body("Where k1 = 1.5 and b = 0.75 control term frequency saturation and document length normalization.")

    add_heading_2("C. Dense Vector Cosine Similarity")
    add_body("Cosine_Similarity(A, B) = (A • B) / (||A|| * ||B||)", "Cosine Formula: ")
    add_body("Measures semantic angle between 768-dimensional query and document vectors regardless of length.")

    # 5. Use Cases
    add_heading_1("5. Enterprise Use Cases & Applications")
    add_body("Clause-by-clause contract comparison, risk redlining, and non-disclosure agreement auditing.", "1. Legal Counsel & Compliance: ")
    add_body("Natural language SQL querying over local database files and automated PowerPoint pitch deck generation.", "2. Financial Auditing & Analysis: ")
    add_body("Air-gapped HIPAA-compliant patient medical document indexing and clinical research synthesis.", "3. Healthcare & Clinical Research: ")
    add_body("Offline codebase search across Python/JS files with automated Mermaid.js architecture mind maps.", "4. Software Engineering & Architecture: ")

    # 6. Future Roadmap
    add_heading_1("6. Future Roadmap & Strategic Next Steps")
    add_body("Integrate local LLaVA / Llama 3.2-Vision models directly into PyInstaller bundle for full image understanding without Tesseract.", "1. Native Vision LLM Integration: ")
    add_body("Add BGE-Reranker-Large cross-encoder model for 2-stage retrieval re-ranking.", "2. Cross-Encoder Re-ranking: ")
    add_body("Implement sub-300ms WebRTC voice stream processing for fluid real-time voice conversations.", "3. Ultra-Low Latency Voice Loop: ")
    add_body("Package the application inside an Electron.js desktop application wrapper for native Windows/macOS window management.", "4. Native Electron Desktop App: ")

    doc.save(DOCX_PATH)
    print(f"[SUCCESS] Word documentation created: {DOCX_PATH}")


def generate_pdf_documentation():
    if fitz is None:
        print("pymupdf not installed, skipping pdf generation.")
        return

    doc = fitz.open()

    # Read docx or construct clean PDF pages directly with PyMuPDF
    page = doc.new_page(width=595, height=842) # A4
    margin = 50
    y = 50

    # Draw Title
    page.insert_text((margin, y), "ORBIT PLATFORM DOCUMENTATION", fontsize=20, fontname="helv", color=(0.24, 0.25, 0.94))
    y += 25
    page.insert_text((margin, y), "Offline Multimodal RAG Engine — Full Technical Specification", fontsize=11, fontname="helv", color=(0.5, 0.2, 0.8))
    y += 30

    sections = [
        ("1. EXECUTIVE OVERVIEW", "Orbit is a 100% offline Multimodal RAG Platform providing enterprise privacy, air-gapped document intelligence, and zero cloud data leakage."),
        ("2. KEY FEATURES", "• Support for PDF, Word, PowerPoint, Excel, Images, Audio, Video, SQLite DBs.\n• 3-Column Orbit Dashboard UI.\n• Hybrid BM25 Keyword + Dense Vector Search.\n• Agentic Multi-Hop Query Decomposition & GraphRAG.\n• Self-RAG Reflective Critic Agent."),
        ("3. TECHNICAL STACK", "• Frontend: Streamlit 1.39+ with CSS Glassmorphism\n• Local LLM: Ollama llama3.2:3b\n• Embedding Model: nomic-embed-text (768-dim)\n• Vector Store: ChromaDB Persistent HNSW Client\n• Search Engine: Rank-BM25 (BM25Okapi)\n• Transcription: faster-whisper (CTranslate2)\n• Executable: PyInstaller 6.22 Standalone EXE"),
        ("4. ALGORITHMS USED", "• Reciprocal Rank Fusion (RRF): Score = 1 / (60 + vec_rank) + 1 / (60 + bm25_rank)\n• Okapi BM25 Keyword Scoring (k1=1.5, b=0.75)\n• Dense Vector Cosine Similarity\n• Overlapping Recursive Character Window Chunking"),
        ("5. FUTURE IMPLEMENTATIONS", "• Local LLaVA Multimodal Vision LLM Integration\n• Cross-Encoder BGE-Reranker-Large\n• Electron.js Native Desktop App Packaging\n• Sub-300ms WebRTC Real-Time Voice Mode"),
    ]

    for title, content in sections:
        page.insert_text((margin, y), title, fontsize=12, fontname="helv", color=(0.24, 0.25, 0.94))
        y += 18
        for line in content.split("\n"):
            page.insert_text((margin + 10, y), line, fontsize=9.5, fontname="helv", color=(0.2, 0.2, 0.2))
            y += 14
        y += 10

    doc.save(PDF_PATH)
    doc.close()
    print(f"[SUCCESS] PDF documentation created: {PDF_PATH}")


if __name__ == "__main__":
    generate_docx_documentation()
    generate_pdf_documentation()
