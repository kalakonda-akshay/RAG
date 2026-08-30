"""
Generates the comprehensive, highly styled Orbit User & Installation Guide PDF and DOCX documents with embedded screenshots.
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak, HRFlowable, KeepTogether
)
from reportlab.pdfgen import canvas

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "data", "reports")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Screenshot image paths
IMG_CHAT = os.path.join(BASE_DIR, "offlinerag", "public", "screenshots", "chat_citations.png")
IMG_VOICE = os.path.join(BASE_DIR, "offlinerag", "public", "screenshots", "voice_call.png")
IMG_QUIZ = os.path.join(BASE_DIR, "offlinerag", "public", "screenshots", "quiz_workspace.png")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

class NumberedCanvas(canvas.Canvas):
    """Canvas for adding page numbers and running footer."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748B"))
        
        # Header
        self.drawString(36, 760, "ORBIT — Offline Multimodal RAG Assistant | User & Installation Guide")
        self.setStrokeColor(colors.HexColor("#E2E8F0"))
        self.setLineWidth(0.5)
        self.line(36, 752, 576, 752)
        
        # Footer
        self.line(36, 45, 576, 45)
        self.drawString(36, 32, "Confidential & Private — 100% Air-Gapped Intelligence")
        self.drawRightString(576, 32, f"Page {self._pageNumber} of {page_count}")
        self.restoreState()

def generate_pdf():
    pdf_path = os.path.join(OUTPUT_DIR, "Orbit_Complete_User_Guide.pdf")
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=54,
        bottomMargin=54
    )
    story = []
    styles = getSampleStyleSheet()

    # Custom typography styles
    title_style = ParagraphStyle(
        'MainTitle',
        parent=styles['Heading1'],
        fontSize=22,
        leading=26,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#4F46E5'),
        alignment=1,
        spaceAfter=6
    )

    subtitle_style = ParagraphStyle(
        'MainSubTitle',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        fontName='Helvetica',
        textColor=colors.HexColor('#64748B'),
        alignment=1,
        spaceAfter=15
    )

    h1_style = ParagraphStyle(
        'Heading1_Custom',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'Heading2_Custom',
        parent=styles['Heading3'],
        fontSize=11.5,
        leading=15,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#4F46E5'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'Body_Custom',
        parent=styles['BodyText'],
        fontSize=9.5,
        leading=13.5,
        fontName='Helvetica',
        textColor=colors.HexColor('#334155'),
        spaceAfter=6
    )

    callout_style = ParagraphStyle(
        'CalloutText',
        parent=body_style,
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#1E1B4B')
    )

    # 1. Header Banner
    story.append(Paragraph("ORBIT — OFFLINE MULTIMODAL RAG ASSISTANT", title_style))
    story.append(Paragraph("COMPLETE USER MANUAL, INSTALLATION GUIDE & APP WALKTHROUGH", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#4F46E5'), spaceAfter=14))

    # 2. Executive Overview
    story.append(Paragraph("1. Executive Overview & Mission", h1_style))
    story.append(Paragraph(
        "<b>Orbit</b> is an enterprise-grade, local-first Retrieval-Augmented Generation (RAG) assistant built for complete data privacy. Powered by dense vector embeddings (<i>nomic-embed-text</i>), sparse keyword reranking (<i>Rank-BM25</i>), and local LLMs (<i>llama3.2:3b</i>), Orbit analyzes your documents, spreadsheets, images, audio, video, and SQL databases without transferring any information to cloud servers.", body_style
    ))

    # Metric Table
    table_data = [
        [Paragraph("<b>Platform Name</b>", body_style), Paragraph("Orbit (Offline Multimodal RAG Assistant)", body_style)],
        [Paragraph("<b>Live Web Portal</b>", body_style), Paragraph("<font color='#4F46E5'><u>https://offlinerag-psi.vercel.app</u></font>", body_style)],
        [Paragraph("<b>Direct Download URL</b>", body_style), Paragraph("<font color='#4F46E5'><u>https://github.com/kalakonda-akshay/RAG/releases/latest/download/OfflineRAGAssistant_Setup.zip</u></font>", body_style)],
        [Paragraph("<b>Supported Formats</b>", body_style), Paragraph("<b>29 File Extensions</b> (PDF, DOCX, XLSX, PPTX, MP3, MP4, PNG, SQLite)", body_style)],
        [Paragraph("<b>Local AI Engine</b>", body_style), Paragraph("Ollama (llama3.2:3b + nomic-embed-text 768-dim)", body_style)]
    ]
    t_overview = Table(table_data, colWidths=[140, 380])
    t_overview.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#F1F5F9')),
        ('BACKGROUND', (1,0), (1,-1), colors.HexColor('#F8FAFC')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('PADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(t_overview)
    story.append(Spacer(1, 12))

    # 3. Installation Guide
    story.append(Paragraph("2. Download & Installation Guide", h1_style))
    
    # Callout Box
    callout_data = [[Paragraph("<b>🔒 Air-Gapped Privacy Guarantee:</b> Orbit runs 100% locally on your computer CPU/GPU. No internet connection is required after initial model download.", callout_style)]]
    t_callout = Table(callout_data, colWidths=[520])
    t_callout.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#EEF2FF')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#6366F1')),
        ('PADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t_callout)
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>Step 01: Download Setup Package</b>", h2_style))
    story.append(Paragraph("Download <i>OfflineRAGAssistant_Setup.zip</i> (1.9 GB) from <font color='#4F46E5'>https://offlinerag-psi.vercel.app</font> or GitHub Releases.", body_style))

    story.append(Paragraph("<b>Step 02: Extract & Environment Setup</b>", h2_style))
    story.append(Paragraph("Extract the zip file to your folder (e.g. <i>C:\\OrbitAssistant</i>) and run the automated setup wizard in terminal:<br/><code>python launcher/first_run_setup.py</code>", body_style))

    story.append(Paragraph("<b>Step 03: Pull Local Ollama LLM Models</b>", h2_style))
    story.append(Paragraph("Run the following terminal commands to pull local models:<br/><code>ollama pull llama3.2:3b</code><br/><code>ollama pull nomic-embed-text</code>", body_style))

    story.append(Paragraph("<b>Step 04: Launch Web Interface</b>", h2_style))
    story.append(Paragraph("Start the application server: <code>python launcher/run_app.py</code><br/>Access URL: <b>http://localhost:8501</b>", body_style))
    story.append(Spacer(1, 12))

    # 4. Opening the App & Walkthrough
    story.append(Paragraph("3. Opening the App — Layout & Interface Walkthrough", h1_style))
    story.append(Paragraph(
        "When you launch Orbit at <b>http://localhost:8501</b>, you enter a modern 3-column dark-mode workspace. Here is a walkthrough of what you see after opening the app:", body_style
    ))

    rail_data = [
        [Paragraph("<b>Workspace Component</b>", body_style), Paragraph("<b>Description & Features Available After Opening</b>", body_style)],
        [
            Paragraph("<b>Left Knowledge Rail</b><br/>(Workspace & File Manager)", body_style),
            Paragraph("• <b>Brand Header</b>: Orbit logo and version indicator.<br/>• <b>Workspace Switcher</b>: Select between <i>documents</i>, <i>research</i>, <i>finance</i>, and <i>engineering</i>.<br/>• <b>Indexed File Cards</b>: Displays loaded files with type icons (PDF, XLSX, DOCX).<br/>• <b>Upload Dropzone</b>: Drag & drop any of the 29 supported file types.<br/>• <b>AI Personas</b>: Switch between <i>Technical Analyst</i>, <i>Tutor</i>, and <i>Legal Officer</i>.<br/>• <b>Export Button</b>: 1-Click PowerPoint deck export (<i>Executive_Deck.pptx</i>).", body_style)
        ],
        [
            Paragraph("<b>Center Q&A Thread</b><br/>(Feature Dashboard & Chat)", body_style),
            Paragraph("• <b>Scope Banner</b>: Shows active scope (e.g. <i>scoped to research (1 file)</i>).<br/>• <b>12 Feature Tabs</b>: Quick switcher for Chat, Voice Call, Web Archiver, Mind Map, Topics, Translate, Battle, Quiz, SQL DB, PII, Compare, Reports.<br/>• <b>Message Stream</b>: Displays grounded AI answers with inline <b>[1]</b>, <b>[2]</b> citation chips.<br/>• <b>Interactive Source Pills</b>: Clickable source buttons (e.g. <i>📄 [1] Company Financials Q1 (p.4)</i>).<br/>• <b>Prompt Input Bar</b>: Type grounded questions over documents.", body_style)
        ],
        [
            Paragraph("<b>Right Grounding Rail</b><br/>(Source Inspection)", body_style),
            Paragraph("• <b>Grounding Gauge</b>: Shows citation verification rating (e.g. <i>0.91 high</i>).<br/>• <b>Relevance Cards</b>: Shows retrieved text passages with cosine similarity scores.<br/>• <b>Highlighted Text</b>: Displays exact matched text with <code>&lt;mark&gt;</code> highlights.<br/>• <b>Page & Chunk Markers</b>: Shows exact page numbers and chunk boundaries.", body_style)
        ]
    ]
    t_rails = Table(rail_data, colWidths=[150, 370])
    t_rails.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (1,0), colors.HexColor('#4F46E5')),
        ('TEXTCOLOR', (0,0), (1,0), colors.white),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('BACKGROUND', (0,1), (0,-1), colors.HexColor('#F1F5F9')),
        ('BACKGROUND', (1,1), (1,-1), colors.HexColor('#FFFFFF')),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_rails)
    story.append(Spacer(1, 14))

    # 5. Screenshots Tour
    story.append(Paragraph("4. Live Interface & Screenshot Tour", h1_style))

    if os.path.exists(IMG_CHAT):
        story.append(Paragraph("<b>Figure 1: Grounded Q&A Thread & Interactive Source Pills</b>", h2_style))
        story.append(Paragraph("Displays the main chat thread, inline citation badges <b>[1]</b>, <b>[2]</b>, interactive document pills, and prompt input box:", body_style))
        story.append(RLImage(IMG_CHAT, width=520, height=292))
        story.append(Spacer(1, 12))

    if os.path.exists(IMG_VOICE):
        story.append(Paragraph("<b>Figure 2: Real-Time Voice Call Mode</b>", h2_style))
        story.append(Paragraph("Displays the hands-free microphone audio recording interface powered by local Whisper speech-to-text:", body_style))
        story.append(RLImage(IMG_VOICE, width=520, height=292))
        story.append(Spacer(1, 12))

    if os.path.exists(IMG_QUIZ):
        story.append(Paragraph("<b>Figure 3: Auto-Generated Document Quiz Workspace</b>", h2_style))
        story.append(Paragraph("Displays practice test questions auto-generated from indexed document contents:", body_style))
        story.append(RLImage(IMG_QUIZ, width=520, height=292))
        story.append(Spacer(1, 14))

    # 6. Complete Feature Reference
    story.append(Paragraph("5. Complete Feature Reference (12 Feature Tabs)", h1_style))
    features = [
        ("💬 Grounded Q&A Chat", "Dense vector search + Rank-BM25 hybrid retrieval + Self-RAG hallucination verification."),
        ("🎙️ Voice Call Mode", "Real-time microphone recording and local Whisper speech transcription."),
        ("🌐 Web Archiver", "Parses and indexes offline web pages and documentation HTML."),
        ("🗺️ Mind Map Generator", "Generates visual Mermaid.js mind maps from document concepts."),
        ("🧩 Topic Cluster Discovery", "Performs K-Means semantic clustering across document text."),
        ("🌐 Multilingual Translator", "Translates document passages into 10+ target languages offline."),
        ("⚔️ Model Battle Mode", "Compares responses side-by-side across multiple local LLMs."),
        ("🎓 Auto Document Quiz", "Generates practice test questions with answer keys."),
        ("🗄️ Natural Language SQL Engine", "Queries SQLite (.db/.sqlite) tables using auto-generated SQL."),
        ("🛡️ PII Redactor & PDF Exporter", "Redacts SSNs, credit cards, emails, and exports sanitized PDFs."),
        ("🔍 Document Comparator", "Side-by-side comparison of two documents highlighting diffs."),
        ("📊 Report & PPTX Exporter", "1-click PowerPoint presentations (.pptx) and executive brief summaries.")
    ]
    for title, desc in features:
        story.append(Paragraph(f"<b>{title}</b>: {desc}", body_style))

    story.append(Spacer(1, 12))

    # 7. 29 File Formats
    story.append(Paragraph("6. Supported File Formats (29 Extensions)", h1_style))
    story.append(Paragraph(
        "• <b>Documents & Presentations (5)</b>: .pdf, .docx, .pptx, .ppt, .md<br/>"
        "• <b>Spreadsheets & Tables (3)</b>: .xlsx, .xls, .csv<br/>"
        "• <b>Images & OCR (7)</b>: .png, .jpg, .jpeg, .webp, .bmp, .tiff, .gif<br/>"
        "• <b>Audio Recordings (3)</b>: .wav, .mp3, .m4a<br/>"
        "• <b>Video Files (4)</b>: .mp4, .mkv, .mov, .avi<br/>"
        "• <b>Code & Text Markup (7)</b>: .txt, .py, .js, .json, .html, .xml, .sql<br/>"
        "• <b>Relational DBs (2)</b>: .sqlite, .db",
        body_style
    ))

    story.append(Spacer(1, 12))

    # 8. Algorithm Architecture
    story.append(Paragraph("7. Algorithmic Architecture & RAG Pipeline", h1_style))
    story.append(Paragraph(
        "1. <b>Multimodal Ingestion</b>: Extracts text, tables, image OCR, and audio voice files.<br/>"
        "2. <b>Recursive Chunker</b>: Overlapping sliding window (500 tokens / 50 overlap).<br/>"
        "3. <b>Dense Embeddings</b>: nomic-embed-text 768-dimensional dense vectors.<br/>"
        "4. <b>Hybrid RRF Search</b>: Dense Cosine Similarity + Sparse Rank-BM25 Reciprocal Rank Fusion.<br/>"
        "5. <b>Agentic Multi-Hop RAG</b>: Recursive sub-query splitting.<br/>"
        "6. <b>Local LLM Generator & Self-RAG Critic</b>: Ollama llama3.2:3b fact audit.<br/>"
        "7. <b>GraphRAG</b>: Entity-relationship triple extraction (Source -> Relation -> Target).",
        body_style
    ))

    # Build PDF with custom NumberedCanvas
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Saved Styled PDF User Guide to: {pdf_path}")

def generate_docx():
    docx_path = os.path.join(OUTPUT_DIR, "Orbit_Complete_User_Guide.docx")
    doc = Document()

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("ORBIT — OFFLINE MULTIMODAL RAG ASSISTANT\nCOMPLETE USER & INSTALLATION GUIDE")
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(79, 70, 229)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("100% Air-Gapped Intelligence · Installation · Post-Launch Walkthrough · Feature Guide")
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Executive Summary
    doc.add_heading("1. Executive Overview & Mission", level=1)
    doc.add_paragraph(
        "Orbit is an advanced, local-first Retrieval-Augmented Generation (RAG) platform designed to deliver enterprise-grade document intelligence without sending a single byte of data to external cloud servers. Powered by local vector search, sparse Rank-BM25 recency reranking, and Ollama local LLMs (llama3.2:3b), Orbit ensures complete privacy, zero API costs, and air-gapped security."
    )

    # 2. Installation Guide
    doc.add_heading("2. Download & Installation Step-by-Step Guide", level=1)
    doc.add_paragraph(
        "• Download Setup Package: https://github.com/kalakonda-akshay/RAG/releases/latest/download/OfflineRAGAssistant_Setup.zip\n"
        "• Web Portal: https://offlinerag-psi.vercel.app\n"
        "• Environment Setup: python launcher/first_run_setup.py\n"
        "• Pull Local LLM Models: ollama pull llama3.2:3b && ollama pull nomic-embed-text\n"
        "• Launch Application: python launcher/run_app.py (http://localhost:8501)"
    )

    # 3. Opening the App Walkthrough
    doc.add_heading("3. Opening the App — Layout & Interface Walkthrough", level=1)
    doc.add_paragraph(
        "When you open Orbit at http://localhost:8501, you are presented with a 3-column interactive layout:\n\n"
        "1. Left Knowledge Rail: Brand header, workspace switcher (documents, research, finance, engineering), indexed file cards, upload dropzone, AI persona switcher, export PPTX button.\n"
        "2. Center Q&A Thread: Scope banner, 12 feature navigation tabs, message stream with [1], [2] citations, interactive source pills, raw text chunk inspector, prompt bar.\n"
        "3. Right Grounding Rail: Grounding confidence meter (0.91 high), cited document relevance cards, <mark> text highlights, page and chunk markers."
    )

    # 4. Screenshots Tour
    doc.add_heading("4. Visual Screenshot Tour", level=1)
    if os.path.exists(IMG_CHAT):
        doc.add_heading("Figure 1: Grounded Q&A Thread & Interactive Citations", level=2)
        doc.add_picture(IMG_CHAT, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists(IMG_VOICE):
        doc.add_heading("Figure 2: Real-Time Voice Call Mode", level=2)
        doc.add_picture(IMG_VOICE, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists(IMG_QUIZ):
        doc.add_heading("Figure 3: Auto-Generated Document Quiz Workspace", level=2)
        doc.add_picture(IMG_QUIZ, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. Complete Feature Reference
    doc.add_heading("5. Complete Feature & Workspace Reference (12 Tabs)", level=1)
    features = [
        ("💬 Grounded Q&A Chat", "Dense vector search + Rank-BM25 hybrid retrieval + Self-RAG verification."),
        ("🎙️ Voice Call Mode", "Hands-free voice interface using microphone input and local Whisper speech transcription."),
        ("🌐 Web Archiver", "Parses and indexes offline web pages, documentation HTML, and articles."),
        ("🗺️ Mind Map Generator", "Generates interactive visual Mermaid.js mind maps from uploaded document concepts."),
        ("🧩 Topic Cluster Discovery", "Performs K-Means semantic clustering to discover hidden themes across files."),
        ("🌐 Multilingual Translator", "Translates document passages into 10+ target languages completely offline."),
        ("⚔️ Model Battle Mode", "Compares responses side-by-side across multiple local LLMs."),
        ("🎓 Auto Document Quiz", "Generates practice test questions with answer keys from document text."),
        ("🗄️ Natural Language SQL Engine", "Queries SQLite (.db/.sqlite) tables using auto-generated SQL code."),
        ("🛡️ PII Redactor & PDF Exporter", "Detects and redacts SSNs, credit cards, emails, and phone numbers, exporting sanitized PDFs."),
        ("🔍 Document Comparator", "Compares two documents side-by-side highlighting overlaps, diffs, and contradictions."),
        ("📊 Report & PPTX Exporter", "Generates 1-click PowerPoint presentations (.pptx) and executive brief summaries.")
    ]
    for title, desc in features:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{title}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(79, 70, 229)
        p.add_run(desc)

    doc.save(docx_path)
    print(f"Saved DOCX User Guide to: {docx_path}")

if __name__ == "__main__":
    generate_pdf()
    generate_docx()
